from django.shortcuts import render, redirect, get_object_or_404
from django.conf import settings
from azure.storage.blob import BlobServiceClient
from docx import Document
from urllib.parse import urlparse, unquote
from io import BytesIO
import PyPDF2
from .models import Documento
from .utils import (
    detectar_idioma,
    analizar_sentimiento,
    contar_palabras,
    contar_oraciones,
    puntuacion_documento,
    obtener_codigo_languagetool,
    evaluar_cohesion_local,
    detectar_errores,
    generar_preguntas_variadas,
    dashboard_metricas,
    generar_links_palabras_clave,
    traducir_texto,
    extraer_palabras_clave,
    obtener_nombre_idioma_de_traduccion
)


def analizar_documento(request, documento_id):
    documento = get_object_or_404(Documento, id=documento_id, usuario=request.user)

    blob_service_client = BlobServiceClient.from_connection_string(settings.AZURE_STORAGE_CONNECTION_STRING)
    ruta_blob = urlparse(str(documento.archivo)).path
    blob_name = unquote(ruta_blob).split(f'/{settings.AZURE_STORAGE_CONTAINER_NAME}/')[-1]

    try:
        blob_client = blob_service_client.get_blob_client(container=settings.AZURE_STORAGE_CONTAINER_NAME, blob=blob_name)
        stream = blob_client.download_blob()
        contenido = stream.readall()

        extension = documento.archivo.name.split('.')[-1].lower()

        if extension in ["txt", "csv", "log"]:
            texto_original = contenido.decode("utf-8", errors="ignore")
        elif extension == "pdf":
            pdf_reader = PyPDF2.PdfReader(BytesIO(contenido))
            texto_original = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        elif extension == "docx":
            doc = Document(BytesIO(contenido))
            texto_original = "\n".join([p.text for p in doc.paragraphs])
        else:
            return render(request, 'error.html', {'mensaje': f"⚠ Formato de archivo '{extension}' no soportado."})

    except Exception as e:
        return render(request, 'error.html', {'mensaje': f"Error al descargar desde Azure: {str(e)}"})

    idioma_detectado = detectar_idioma(texto_original)
    analizar_sent_doc = analizar_sentimiento(texto_original)
    palabras_contadas_original = contar_palabras(texto_original)
    oraciones_contadas_original = contar_oraciones(texto_original)
    resultado = puntuacion_documento(texto_original)
    mapeo_idioma_az_langtool = obtener_codigo_languagetool(idioma_detectado)
    nivel_cohesion = evaluar_cohesion_local(texto_original)
    correcciones = detectar_errores(texto_original, mapeo_idioma_az_langtool)
    preguntas_vf = generar_preguntas_variadas(texto_original)
    diagrama = dashboard_metricas(texto_original)
    links_bibliograficos = generar_links_palabras_clave(texto_original, mapeo_idioma_az_langtool)
    return render(request, 'documentos/analisis.html', {
        'documento': documento,
        'texto_original': texto_original,
        'idioma_detectado': idioma_detectado,
        'sentimiento': analizar_sent_doc,
        'total_palabras': palabras_contadas_original,
        'total_oraciones': oraciones_contadas_original,
        'total_puntuacion_lectora': resultado,
        'correcciones_detectadas': correcciones,
        'cohesion': nivel_cohesion,
        'preguntas_vyf': preguntas_vf,
        **diagrama,
        'enlaces_clave_google': links_bibliograficos
    })


def traducir_documento(request, documento_id):
    documento = get_object_or_404(Documento, id=documento_id, usuario=request.user)

    if request.method == "POST":
        idioma_destino = request.POST.get("idioma_destino")

        blob_service_client = BlobServiceClient.from_connection_string(settings.AZURE_STORAGE_CONNECTION_STRING)
        ruta_blob = urlparse(str(documento.archivo)).path
        blob_name = unquote(ruta_blob).split(f'/{settings.AZURE_STORAGE_CONTAINER_NAME}/')[-1]

        try:
            blob_client = blob_service_client.get_blob_client(container=settings.AZURE_STORAGE_CONTAINER_NAME, blob=blob_name)
            stream = blob_client.download_blob()
            contenido = stream.readall()

            extension = documento.archivo.name.split('.')[-1].lower()

            if extension == "txt":
                texto_original = contenido.decode("utf-8", errors="ignore")
            elif extension == "pdf":
                pdf_reader = PyPDF2.PdfReader(BytesIO(contenido))
                texto_original = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
            elif extension == "docx":
                doc = Document(BytesIO(contenido))
                texto_original = "\n".join([p.text for p in doc.paragraphs])
            else:
                return render(request, 'error.html', {'mensaje': f"⚠ Formato de archivo '{extension}' no soportado."})
            
            texto_traducido = traducir_texto(texto_original, idioma_destino)          
            idioma_detectado = detectar_idioma(texto_traducido)
            mapeo_idioma_azure = obtener_codigo_languagetool(idioma_detectado)
            palabras_clave = extraer_palabras_clave(texto_traducido, mapeo_idioma_azure)
            idioma_seleccionado_usuario = obtener_nombre_idioma_de_traduccion(idioma_destino)
            palabras_contadas_traducidas = contar_palabras(texto_traducido)
            oraciones_contadas_traducidas = contar_oraciones(texto_traducido)

            return render(request, 'documentos/traduccion_final.html', {
                'documento': documento,
                'texto_original': texto_original,
                'texto_traducido': texto_traducido,
                'idioma_destino': idioma_destino,
                'palabras_clave': palabras_clave,
                'idioma_seleccionado_user': idioma_seleccionado_usuario,
                'total_palabras_traducidas': palabras_contadas_traducidas,
                'total_oraciones_traducidas': oraciones_contadas_traducidas
            })

        except Exception as e:
            return render(request, 'error.html', {'mensaje': f"Error al descargar desde Azure: {str(e)}"})

    return redirect('lista_documentos')
