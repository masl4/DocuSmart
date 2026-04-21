import os
import io
import re
import time
import base64
import random
import datetime
from urllib.parse import urlparse, unquote
from collections import Counter

import requests
import pdfplumber
import PyPDF2
import pandas as pd
import language_tool_python
import textstat
import spacy
from io import BytesIO
from docx import Document
from wordcloud import WordCloud
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet

from azure.storage.blob import BlobServiceClient

from django.conf import settings
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required

from .forms import DocumentoForm
from .models import Documento


# ========================
# Funciones auxiliares
# ========================

def azure_post(url, headers, body):
    try:
        response = requests.post(url, headers=headers, json=body)
        response.raise_for_status()
        return response.json()
    except Exception as e:
        return {"error": str(e)}

def azure_headers(service="text"):
    headers = {
        "Content-Type": "application/json",
        "Ocp-Apim-Subscription-Key": (
            settings.AZURE_TEXT_ANALYTICS_KEY if service in ["text", "textanalytics"] else settings.AZURE_TRANSLATOR_KEY
        ),
    }
    if service == "translator":
        headers["Ocp-Apim-Subscription-Region"] = settings.AZURE_TRANSLATOR_REGION
    return headers

# ========================
# Funciones traduccion.html y analisis.html
# ========================


@login_required(login_url='/usuarios/login/')
def subir_documento(request):
    if request.method == 'POST':
        form = DocumentoForm(request.POST, request.FILES)
        if form.is_valid():
            archivo = request.FILES['archivo']

            #  nombre único con timestamp para que no haya dobles
            timestamp = int(time.time())
            nombre_archivo = f"usuario_{request.user.id}/{timestamp}_{archivo.name}"

            # Subimos el archivo a Azure Blob Storage
            url_archivo = subir_archivo_a_azure(archivo, archivo.name, request.user.id)

            if url_archivo:
                print(f"📂 Archivo subido a Azure: {url_archivo}")
                
                # Guardamos en la base de datos
                documento = form.save(commit=False)
                documento.usuario = request.user
                documento.archivo = url_archivo  # Guardamos la URL
                documento.save()

                return redirect('lista_documentos')
            else:
                return render(request, 'error.html', {'mensaje': 'Error al subir el archivo a Azure.'})

    else:
        form = DocumentoForm()
    
    return render(request, 'documentos/subir.html', {'form': form})

def eliminar_documento(request, documento_id):
    documento = get_object_or_404(Documento, id=documento_id, usuario=request.user)

    try:
        # Conexion a Azure Blob Storage
        blob_service_client = BlobServiceClient.from_connection_string(settings.AZURE_STORAGE_CONNECTION_STRING)

        # Extraigo la ruta completa del blob desde la URL
        archivo_url = str(documento.archivo)
        contenedor = settings.AZURE_STORAGE_CONTAINER_NAME

        # Extraer solo el path del blob (después del nombre del contenedor)
        ruta_blob = urlparse(archivo_url).path
        blob_path = unquote(ruta_blob).split(f'/{contenedor}/')[-1]


        # Eliminar el archivo del contenedor de Azure
        blob_client = blob_service_client.get_blob_client(container=contenedor, blob=blob_path)
        blob_client.delete_blob()

        # Eliminar el documento de la base de datos
        documento.delete()

        return JsonResponse({'success': True})

    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

def lista_documentos(request):
    documentos = Documento.objects.filter(usuario=request.user)

    # Extraer solo el nombre del archivo para cada documento
    for documento in documentos:
        ruta_blob = urlparse(str(documento.archivo)).path
        nombre_completo = os.path.basename(ruta_blob)
        # Quitar timestamp para que no moleste luego en la lista
        if "_" in nombre_completo:
            nombre_sin_timestamp = "_".join(nombre_completo.split("_")[1:])
            documento.nombre_archivo = nombre_sin_timestamp
        else:
            documento.nombre_archivo = nombre_completo
    return render(request, 'documentos/lista.html', {'documentos': documentos})

def extraer_texto(archivo):
    ext = archivo.name.lower().split('.')[-1]

    if ext == "pdf":
        with pdfplumber.open(archivo) as pdf:
            return "\n".join(filter(None, [page.extract_text() for page in pdf.pages]))
    elif ext == "docx":
        return "\n".join([p.text for p in Document(archivo).paragraphs])
    elif ext == "txt":
        return archivo.read().decode("utf-8")
    
    return None


def contar_palabras(texto):
    return len(texto.split())


def contar_oraciones(texto):
    return len(texto.split("."))


def subir_archivo_a_azure(archivo, nombre_original, user_id):
    try:
        blob_service_client = BlobServiceClient.from_connection_string(settings.AZURE_STORAGE_CONNECTION_STRING)

        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        nombre_archivo = f"usuario_{user_id}/{timestamp}_{nombre_original}"

        blob_client = blob_service_client.get_blob_client(
            container=settings.AZURE_STORAGE_CONTAINER_NAME,
            blob=nombre_archivo
        )

        blob_client.upload_blob(archivo, overwrite=True)

        cuenta = settings.AZURE_STORAGE_CONNECTION_STRING.split(';')[1].split('=')[1]
        return f"https://{cuenta}.blob.core.windows.net/{settings.AZURE_STORAGE_CONTAINER_NAME}/{nombre_archivo}"
    except Exception as e:
        print(f"❌ Error subiendo archivo a Azure: {e}")
        return None


def error_view(request, mensaje):
    return render(request, 'error.html', {'mensaje': mensaje})

# ========================
# Funciones analisis.html
# ========================


def detectar_idioma(texto):
    url = f"{settings.AZURE_TRANSLATOR_ENDPOINT.rstrip('/')}/detect?api-version=3.0"
    headers = azure_headers("translator")
    body = [{"text": texto[:5000]}]

    try:
        result = azure_post(url, headers, body)

        if isinstance(result, list) and result:
            idioma = result[0].get("language")
            if idioma:
                return idioma
        raise ValueError("No se detectó el idioma en la respuesta de Azure.")
    except Exception as e:
        raise RuntimeError(f"Error al detectar idioma: {str(e)}")




def extraer_palabras_clave(texto, idioma_destino):
    url = f"{settings.AZURE_TEXT_ANALYTICS_ENDPOINT.rstrip('/')}/text/analytics/v3.1/keyPhrases"
    body = {
        "documents": [{
            "id": "1",
            "language": idioma_destino,
            "text": texto[:5000]
        }]
    }

    try:
        result = azure_post(url, azure_headers("text"), body)
        documentos = result.get("documents", [])
        if documentos and "keyPhrases" in documentos[0]:
            frases = documentos[0]["keyPhrases"]
            return [{"texto": palabra, "relevancia": i + 1} for i, palabra in enumerate(frases)]
        return []
    except Exception:
        return []




def analizar_sentimiento(texto):
    url = f"{settings.AZURE_TEXT_ANALYTICS_ENDPOINT}/text/analytics/v3.0/sentiment"
    headers = azure_headers("textanalytics")
    
    documentos = [{
        "id": "1",
        "language": "es",
        "text": texto[:5000] 
    }]

    try:
        result = azure_post(url, headers, {"documents": documentos})

        if isinstance(result, dict) and "documents" in result and result["documents"]:
            sentimiento = result["documents"][0].get("sentiment", "indeterminado")
            return sentimiento
        else:
            raise ValueError(f"Respuesta inesperada de Azure: {result}")

    except Exception as e:
        return f"❌ Error al analizar sentimiento: {str(e)}"



def puntuacion_documento(texto):
    reading_score = textstat.flesch_reading_ease(texto)

    if reading_score > 89:
        nivel = "5º Primaria o menos"
    elif 80 <= reading_score <= 89:
        nivel = "6º Primaria"
    elif 70 <= reading_score < 80:
        nivel = "1º-2º ESO"
    elif 60 <= reading_score < 70:
        nivel = "3º-4º ESO"
    elif 50 <= reading_score < 60:
        nivel = "1º-2º Bachillerato"
    elif 30 <= reading_score < 50:
        nivel = "Universidad Básica"
    elif 0 <= reading_score < 30:
        nivel = "Universidad Avanzada"
    else:
        nivel = "Nivel no determinado"

    return {
        "reading_score": round(reading_score, 2),
        "nivel_educativo": nivel
    }


def obtener_codigo_languagetool(nombre_idioma):
    mapa_idiomas = {
        'Arabic': 'ar',
        'Catalan': 'ca',
        'Chinese Simplified': 'zh',
        'Chinese Traditional': 'zh-TW',
        'Czech': 'cs',
        'Danish': 'da',
        'Dutch': 'nl',
        'English': 'en-US',
        'Estonian': 'et',
        'Finnish': 'fi',
        'French': 'fr',
        'Galician': 'gl',
        'German': 'de',
        'Greek': 'el',
        'Hebrew': 'he',
        'Hindi': 'hi',
        'Hungarian': 'hu',
        'Indonesian': 'id',
        'Irish': 'ga',
        'Italian': 'it',
        'Japanese': 'ja',
        'Korean': 'ko',
        'Lithuanian': 'lt',
        'Malay': 'ms',
        'Norwegian': 'nb',
        'Polish': 'pl',
        'Portuguese': 'pt',
        'Romanian': 'ro',
        'Russian': 'ru',
        'Slovak': 'sk',
        'Slovenian': 'sl',
        'Spanish': 'es',
        'Swedish': 'sv',
        'Tagalog': 'tl',
        'Thai': 'th',
        'Turkish': 'tr',
        'Ukrainian': 'uk',
        'Vietnamese': 'vi'
    }

    return mapa_idiomas.get(nombre_idioma, 'es')  # 'es' por defecto si no lo encuentra


def detectar_errores(texto, idioma_detectado):

    tool = language_tool_python.LanguageTool(idioma_detectado)
    errores = tool.check(texto)

    sugerencias = []
    for error in errores:
        sugerencias.append({
            "texto_incorrecto": texto[error.offset : error.offset + error.errorLength],
            "mensaje": error.message,
            "reemplazos": error.replacements,
            "contexto": error.context
        })

    return sugerencias

def evaluar_cohesion_local(texto):
    # Dividimos el texto en oraciones
    oraciones = texto.split('.')  
    oraciones = [o.strip() for o in oraciones if o.strip()]  

    if len(oraciones) < 2:
        return "⚠ Texto demasiado corto para evaluar cohesión."

    repeticiones = 0 

    # Comparamos cada oración con la siguiente
    for i in range(len(oraciones) - 1):
        palabras1 = set(oraciones[i].lower().split())
        palabras2 = set(oraciones[i + 1].lower().split())
        interseccion = palabras1 & palabras2  #Sacamos las Palabras comunes

        if len(interseccion) > 2:
            repeticiones += 1  

    # Calcular el nivel de cohesión
    ratio = repeticiones / (len(oraciones) - 1)

    if ratio >= 0.5:
        return "✅ Buena cohesión entre frases."
    elif ratio > 0.2:
        return "⚠ Cohesión aceptable, pero mejorable."
    else:
        return "❌ Baja cohesión. Las frases parecen desconectadas."


def generar_preguntas_variadas(texto, num_preguntas=5):
    prompt = f"""
Genera {num_preguntas} preguntas tipo examen en español que se puedan responder leyendo el siguiente texto. 
Incluye preguntas de respuesta corta y preguntas de verdadero/falso. Devuelve solo una lista numerada sin respuestas.

Texto:
\"\"\"{texto}\"\"\"
"""

    r = requests.post("http://localhost:11434/api/generate", json={
        "model": "mistral",
        "prompt": prompt,
        "stream": False
    })

    respuesta = r.json()["response"]

    preguntas = []
    for linea in respuesta.splitlines():
        if linea.strip() and re.match(r"^\d+\.", linea.strip()):
            texto_pregunta = linea.strip().split(".", 1)[1].strip()

            # Clasificación básica
            if texto_pregunta.lower().startswith(("¿", "qué", "cuál", "dónde", "cómo", "por qué", "quién", "cuándo")):
                tipo = "Corta"

            preguntas.append({
                "pregunta": texto_pregunta,
                "tipo": tipo
            })

    return preguntas



def generar_links_palabras_clave(texto, idioma_destino):
    palabras_clave = extraer_palabras_clave(texto, idioma_destino)
    enlaces = []

    for palabra in palabras_clave:
        query = palabra["texto"].replace(" ", "+")  # Espacios a '+'
        url = f"https://www.google.com/search?q={query}" #Generamos enlaces con las palabras escogidas
        enlace_html = f'<a href="{url}" target="_blank" rel="noopener noreferrer">{palabra["texto"]}</a>'
        enlaces.append({
            "texto": palabra["texto"],
            "relevancia": palabra["relevancia"],
            "link": enlace_html
        })

    return enlaces


def dashboard_metricas(texto):
   
    # Generar nube de palabras
    wordcloud = WordCloud(width=800, height=400, background_color="white").generate(texto)
    buffer = io.BytesIO()
    wordcloud.to_image().save(buffer, format="PNG")
    nube_palabras = base64.b64encode(buffer.getvalue()).decode()

    # Preparar contexto para renderizado
    contexto = {
        "nube_palabras": nube_palabras,
    }

    return contexto


# ========================
# Funciones traduccion.html
# ========================
def seleccionar_idioma(request, documento_id):
    get_object_or_404(Documento, id=documento_id, usuario=request.user)
    return render(request, 'documentos/seleccionar_idioma.html', {'documento_id': documento_id})

import re

def dividir_texto_en_bloques(texto, max_tamano=4900):
    oraciones = re.split(r'(?<=[.!?]) +', texto)
    bloques = []
    bloque_actual = ""

    for oracion in oraciones:
        if len(bloque_actual) + len(oracion) < max_tamano:
            bloque_actual += oracion + " "
        else:
            bloques.append(bloque_actual.strip())
            bloque_actual = oracion + " "
    if bloque_actual:
        bloques.append(bloque_actual.strip())

    return bloques


def traducir_texto(texto, idioma_destino):
    url = f"{settings.AZURE_TRANSLATOR_ENDPOINT.rstrip('/')}/translate?api-version=3.0&to={idioma_destino}"
    headers = azure_headers("translator")
    bloques = dividir_texto_en_bloques(texto)

    traduccion = []

    for bloque in bloques:
        try:
            result = azure_post(url, headers, [{"text": bloque}])

            if isinstance(result, dict) and "error" in result:
                mensaje = result["error"].get("message", "Error desconocido.")
                traduccion.append(f"[❌ Azure error: {mensaje}]")
            elif isinstance(result, list) and result and result[0].get("translations"):
                traduccion.append(result[0]["translations"][0].get("text", ""))
            else:
                traduccion.append("[⚠ Bloque no traducido]")
        except Exception as e:
            traduccion.append(f"[❌ Error en fragmento: {str(e)}]")

    return "\n".join(traduccion).strip()


    
def obtener_nombre_idioma_de_traduccion(codigo_idioma):
    mapa_idiomas = {
        "af": "Afrikáans",
        "sq": "Albanés",
        "am": "Amárico",
        "ar": "Árabe",
        "hy": "Armenio",
        "as": "Asamés",
        "az": "Azerbaiyano (latino)",
        "bn": "Bengalí",
        "ba": "Bashkir",
        "eu": "Vasco",
        "bho": "Bhojpuri",
        "brx": "Bodo",
        "bs": "Bosnio (latino)",
        "bg": "Búlgaro",
        "yue": "Cantonés (tradicional)",
        "ca": "Catalán",
        "lzh": "Chino (clásico)",
        "zh-Hans": "Chino simplificado",
        "zh-Hant": "Chino tradicional",
        "sn": "chiShona",
        "hr": "Croata",
        "cs": "Checo",
        "da": "Danés",
        "prs": "Dari",
        "dv": "Divehi",
        "doi": "Dogri",
        "nl": "Neerlandés",
        "en": "Inglés",
        "et": "Estonio",
        "fo": "Feroés",
        "fj": "Fiyiano",
        "fil": "Filipino",
        "fi": "Finés",
        "fr": "Francés",
        "fr-ca": "Francés (Canadá)",
        "gl": "Gallego",
        "ka": "Georgiano",
        "de": "Alemán",
        "el": "Griego",
        "gu": "Gujarati",
        "ht": "Criollo haitiano",
        "ha": "Hausa",
        "he": "Hebreo",
        "hi": "Hindi",
        "mww": "Hmong Daw (Latín)",
        "hu": "Húngaro",
        "is": "Islandés",
        "ig": "Igbo",
        "id": "Indonesio",
        "ikt": "Inuinnaqtun",
        "iu": "Inuktitut",
        "iu-Latn": "Inuktitut (latino)",
        "ga": "Irlandés",
        "it": "Italiano",
        "ja": "Japonés",
        "kn": "Canarés",
        "ks": "Kashmiri",
        "kk": "Kazajo",
        "km": "Jemer",
        "rw": "Kinyarwanda",
        "tlh-Latn": "Klingon",
        "tlh-Piqd": "Klingon (plqaD)",
        "gom": "Konkani",
        "ko": "Coreano",
        "ku": "Kurdo (central)",
        "kmr": "Kurdo (norte)",
        "ky": "Kirguís (cirílico)",
        "lo": "Lao",
        "lv": "Letón",
        "lt": "Lituano",
        "ln": "Lingala",
        "dsb": "Bajo sorbio",
        "lug": "Luganda",
        "mk": "Macedonio",
        "mai": "Maithili",
        "mg": "Malgache",
        "ms": "Malayo (latino)",
        "ml": "Malayalam",
        "mt": "Maltés",
        "mi": "Maori",
        "mr": "Maratí",
        "mn-Cyrl": "Mongol (cirílico)",
        "mn-Mong": "Mongol (tradicional)",
        "my": "Myanmar",
        "ne": "Nepalí",
        "nb": "Noruego bokmal",
        "nya": "Nyanja",
        "or": "Odia",
        "ps": "Pastún",
        "fa": "Persa",
        "pl": "Polaco",
        "pt": "Portugués (Brasil)",
        "pt-pt": "Portugués (Portugal)",
        "pa": "Punjabi",
        "otq": "Otomí Querétaro",
        "ro": "Rumano",
        "run": "Rundi",
        "ru": "Ruso",
        "sm": "Samoano (latino)",
        "sr-Cyrl": "Serbio (cirílico)",
        "sr-Latn": "Serbio (latino)",
        "st": "Sesotho",
        "nso": "Sotho septentrional",
        "tn": "Setsuana",
        "sd": "Sindhi",
        "si": "Cingalés",
        "sk": "Eslovaco",
        "sl": "Esloveno",
        "so": "Somalí (árabe)",
        "es": "Español",
        "sw": "Swahili (Latín)",
        "sv": "Sueco",
        "ty": "Tahitiano",
        "ta": "Tamil",
        "tt": "Tártaro (Latín)",
        "te": "Telugu",
        "th": "Tailandés",
        "bo": "Tibetano",
        "ti": "Tigriña",
        "to": "Tongano",
        "tr": "Turco",
        "tk": "Turcomano (latino)",
        "uk": "Ucraniano",
        "hsb": "Alto sorbio",
        "ur": "Urdu",
        "ug": "Uigur (árabe)",
        "uz": "Uzbeko (latino)",
        "vi": "Vietnamita",
        "cy": "Galés",
        "xh": "Xhosa",
        "yo": "Yoruba",
        "yua": "Maya Yucateco",
        "zu": "Zulú",
    }

    # Si no encuentra el código, devuelve "Español" como idioma por defecto
    return mapa_idiomas.get(codigo_idioma, "Español")

def descargar_traduccion_pdf(request):
    texto = request.GET.get("texto", "")
    buffer = BytesIO()

    # Para crear documento PDF
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=50, leftMargin=50,
                            topMargin=50, bottomMargin=50)

    styles = getSampleStyleSheet()
    story = []  

    # Añadimos el texto como párrafos
    for parrafo in texto.split("\n"):
        if parrafo.strip():  # Evitar párrafos vacíos
            story.append(Paragraph(parrafo.strip(), styles["Normal"]))
            story.append(Spacer(1, 12))  # Espacio entre párrafos

    doc.build(story)
    buffer.seek(0)

    return HttpResponse(buffer, content_type='application/pdf', headers={
        'Content-Disposition': 'attachment; filename="TraduccionGenerada.pdf"'
    })

# ========================
# Funciones inclusion.html
# ========================

def texto_simplificado(request, documento_id):
    documento = get_object_or_404(Documento, id=documento_id, usuario=request.user)

    # Conectar con Azure Blob Storage usando mis credenciales
    blob_service_client = BlobServiceClient.from_connection_string(settings.AZURE_STORAGE_CONNECTION_STRING)

    # Extraer solo el nombre del archivo en Azure
    ruta_blob = urlparse(str(documento.archivo)).path
    blob_name = unquote(ruta_blob).split(f'/{settings.AZURE_STORAGE_CONTAINER_NAME}/')[-1]

    try:
        blob_client = blob_service_client.get_blob_client(container=settings.AZURE_STORAGE_CONTAINER_NAME, blob=blob_name)
        # Descargamos el archivo
        stream = blob_client.download_blob()
        contenido = stream.readall()

        # Determinamos el tipo de archivo
        extension = documento.archivo.name.split('.')[-1].lower()

        # Leer archivos de texto
        if extension in ["txt", "csv", "log"]:
            texto_original = contenido.decode("utf-8", errors="ignore")
        elif extension == "pdf":
            try:
                pdf_reader = PyPDF2.PdfReader(BytesIO(contenido))
                texto_paginas = []

                for i, page in enumerate(pdf_reader.pages):
                    texto_pagina = page.extract_text()
                    if texto_pagina:
                        texto_paginas.append(texto_pagina)
                    else:
                        print(f"⚠ Página {i+1} no tiene texto extraíble")

                texto_original = "\n".join(texto_paginas)

                if not texto_original.strip():
                    return render(request, 'error.html', {
                        'mensaje': "⚠ No se pudo extraer texto del PDF. Asegúrate de que contiene texto, no solo imágenes."
                    })

            except Exception as e:
                return render(request, 'error.html', {
                    'mensaje': f"❌ Error al procesar el PDF: {str(e)}"
                })

        elif extension == "docx":
            doc = Document(io.BytesIO(contenido))
            texto_original = "\n".join([p.text for p in doc.paragraphs])
        else:
            return render(request, 'error.html', {'mensaje': f"⚠ Formato de archivo '{extension}' no soportado."})

    except Exception as e:
        return render(request, 'error.html', {'mensaje': f"Error al descargar desde Azure: {str(e)}"})

    dislexia = texto_para_accesibilidad(texto_original)

    return render(request, 'documentos/inclusion.html', {
        'documento': documento,
        'texto_original': texto_original,
        'texto_dislexia':dislexia
    })
nlp = spacy.load('es_core_news_sm')

# Diccionario de pictogramas
PICTOGRAMAS = {
    'hola': '👋',
    'mundo': '🌎',
    'amigo': '🧑‍🤝‍🧑',
    'gracias': '🙏',
    'amor': '❤️',
    'escuela': '🏫',
    'felicidad': '😊',
    'tristeza': '😢',
    'familia': '👨‍👩‍👧‍👦',
    'comida': '🍽️',
    'bebida': '🥤',
    'musica': '🎵',
    'deporte': '⚽',
    'libro': '📚',
    'casa': '🏠',
    'viaje': '✈️',
    'naturaleza': '🌳',
    'trabajo': '💼',
    'dinero': '💰',
    'corazon': '❤️',
    'estrella': '⭐',
    'fuego': '🔥',
    'agua': '💧',
    'sol': '☀️',
    'noche': '🌙',
    'ordenador': '💻',
    'telefono': '📱',
    'tiempo': '⌛',
    'reloj': '🕰️',
    'flor': '🌸',
    'animal': '🐶',
    'feliz': '😁',
    'triste': '😭',
    'sorpresa': '😲',
    'amor': '💖',
    'paz': '✌️',
    'fuerza': '💪',
    'fiesta': '🎉',
    'idea': '💡',
    'arte': '🎨',
    'fotografia': '📸',
    'cine': '🎥',
    'cafe': '☕',
    'vino': '🍷',
    'navidad': '🎄',
    'cumpleaños': '🎂',
    'salud': '💊',
    'medicina': '🩺',
    'amistad': '🤗',
    'familia': '👨‍👩‍👧‍👦',
    'misterio': '🕵️',
    'exploracion': '🗺️',
    'avion': '✈️',
    'tren': '🚂',
    'coche': '🚗',
    'bicicleta': '🚲',
    'barco': '🚢',
    'cohete': '🚀',
    'futuro': '🔮',
    'pasado': '⏳',
    'presente': '🕰️',
    'tecnologia': '🖥️',
    'internet': '🌐',
    'robot': '🤖',
    'juego': '🎮',
    'mapa': '🗺️',
    'llave': '🔑',
    'puerta': '🚪',
    'herramienta': '🔧',
    'ciencia': '🔬',
    'matematicas': '➗',
    'historia': '📜',
    'geografia': '🌍',
    'lenguaje': '📝',
    'examen': '📝',
    'victoria': '🏆',
    'perdida': '❌',
    'exito': '🎯',
    'error': '⚠️',
    'apoyo': '🫂',
    'escuchar': '👂',
    'hablar': '🗣️',
    'pensar': '🤔',
    'descanso': '😴',
    'sueño': '💤',
    'belleza': '💄',
    'ropa': '👗',
    'zapatos': '👠',
    'moda': '🕺',
    'boda': '💍',
    'bebé': '👶',
    'superheroe': '🦸',
    'misterio': '🕵️',
    'gato': '🐱',
    'perro': '🐶',
    'leon': '🦁',
    'tigre': '🐯',
    'elefante': '🐘',
    'jirafa': '🦒',
    'caballo': '🐎',
    'pez': '🐟',
    'pajaro': '🐦',
    'mariposa': '🦋',
    'araña': '🕷️',
    'felicidad': '😊',
    'tristeza': '😢',
    'miedo': '😱',
    'enojo': '😡',
    'amor': '❤️',
    'alegria': '😁',
    'sorpresa': '😲',
    'confusion': '😕',
    'frustracion': '😤',
    'paciencia': '🧘',
    'inteligencia': '🧠',
    'fortaleza': '💪',
    'belleza': '💅',
    'amistad': '🤗',
    'generosidad': '🤲',
    'bondad': '💖',
    'esperanza': '🌈',
    'sabiduria': '🦉',
    'misterio': '🕵️',
    'paz': '☮️',
     'inteligencia': '🧠',
    'artificial': '🤖',
    'educación': '📚',
    'herramienta': '🔧',
    'analizar': '📊',
    'datos': '💾',
    'adaptar': '🔁',
    'necesidades': '🎯',
    'estudiante': '🧑‍🎓',
    'enseñanza': '🏫',
    'estudio': '📑',
    'rendimiento': '📈',
    'académico': '🎓',
    'personalizar': '🧩',
    'retroalimentación': '⚡',
    'automatización': '⚙️',
    'tareas': '🗂️',
    'administrativas': '📋',
    'docentes': '👩‍🏫',
    'tiempo': '⌛',
    'apoyo': '🫂',
    'emocional': '❤️',
    'desafíos': '⚠️',
    'éticos': '⚖️',
    'privacidad': '🔒',
    'autonomía': '🚫',
    'protección': '🛡️',
    'personales': '🧾',
    'colaboren': '🤝',
    'desarrolladores': '👨‍💻',
    'tecnologías': '🖥️',
    'ética': '✅',
    'efectiva': '💡',
    'inclusiva': '🌍',
    'equitativa': '⚖️',
        'acceso': '🔑',
    'adulto': '🧑',
    'algoritmo': '📐',
    'alumna': '🧑‍🎓',
    'alumno': '🧑‍🎓',
    'analizar': '📊',
    'anciano': '👴',
    'análisis': '📊',
    'aplicación': '📱',
    'apoyo': '🫂',
    'aprendizaje': '📘',
    'artificial': '🤖',
    'automatización': '⚙️',
    'automatizado': '⚙️',
    'automatizar': '⚙️',
    'base de datos': '🗃️',
    'clase': '🏫',
    'colaboración': '🤝',
    'colaborar': '🤝',
    'colegio': '🏫',
    'computadora': '💻',
    'contenido': '📄',
    'curso': '📅',
    'código': '💻',
    'datos': '💾',
    'desafíos': '⚠️',
    'desarrollador': '👨‍💻',
    'desarrolladora': '👩‍💻',
    'desigualdad': '❌',
    'digital': '💿',
    'digitalización': '🔢',
    'digitalizar': '🧾',
    'docente': '👨‍🏫',
    'educación': '📚',
    'educativo': '📘',
    'eficaz': '✅',
    'eficiencia': '⚡',
    'emociones': '💓',
    'empatía': '🫶',
    'enseñanza': '🏫',
    'equidad': '⚖️',
    'escuela': '🏫',
    'esfuerzo': '💪',
    'estudiante': '🧑‍🎓',
    'estudiar': '📝',
    'estudio': '📖',
    'evaluación': '📊',
    'examen': '📝',
    'familia': '👨‍👩‍👧‍👦',
    'futuro': '🔮',
    'humano': '🧍',
    'igualdad': '⚖️',
    'impacto': '💥',
    'inclusión': '🌍',
    'información': 'ℹ️',
    'informática': '🖥️',
    'inmediata': '⚡',
    'inteligencia': '🧠',
    'inteligencia emocional': '🧠❤️',
    'inteligente': '🧠',
    'interacción': '🗣️',
    'internet': '🌐',
    'lenguaje': '📝',
    'materia': '📚',
    'mejora': '⬆️',
    'moral': '🧭',
    'motivación': '🔥',
    'niña': '👧',
    'niño': '👦',
    'nota': '🧾',
    'nube': '☁️',
    'ordenador': '💻',
    'pantalla': '🖥️',
    'pasado': '⏳',
    'persona': '👤',
    'personalizado': '🧩',
    'personalizar': '🧩',
    'plan': '🗂️',
    'plataforma': '🧩',
    'presente': '🕰️',
    'privacidad': '🔒',
    'procesamiento': '🧠',
    'profesor': '👨‍🏫',
    'profesora': '👩‍🏫',
    'programa': '💽',
    'programar': '👨‍💻',
    'protección': '🛡️',
    'proteger': '🛡️',
    'red': '🕸️',
    'rendimiento': '📈',
    'resultado': '📈',
    'retos': '🚧',
    'retroalimentación': '🔁',
    'revolución': '🌪️',
    'riesgo': '⚠️',
    'robot': '🤖',
    'seguridad': '🔐',
    'servidor': '🖧',
    'sistema': '🧩',
    'social': '🧑‍🤝‍🧑',
    'tareas': '📋',
    'tecnología': '🖥️',
    'tecnológico': '💻',
    'tiempo': '⌛',
    'trabajo': '💼',
    'transformación': '🔄',
    'universidad': '🎓',
    'web': '🌐',
    'ética': '✅',
     'sonriente': '😀',
    'sonrisa abierta': '😃',
    'risa feliz': '😄',
    'dientes sonrisa': '😁',
    'risa fuera de control': '😆',
    'sonrisa nerviosa': '😅',
    'carcajada': '🤣',
    'llorar de risa': '😂',
    'ligera sonrisa': '🙂',
    'al reves': '🙃',
    'derretido': '🫠',
    'guiño': '😉',
    'tierna sonrisa': '😊',
    'angelical': '😇',
    'enamorado': '🥰',
    'ojos corazon': '😍',
    'emocionado': '🤩',
    'beso': '😘',
    'besito': '😗',
    'contento': '☺️',
    'beso con ojos cerrados': '😚',
    'beso feliz': '😙',
    'sonrisa con lagrima': '🥲',
    'sabroso': '😋',
    'sacando lengua': '😛',
    'guiño lengua': '😜',
    'loco divertido': '🤪',
    'risa con lengua': '😝',
    'cara dinero': '🤑',
    'abrazo': '🤗',
    'mano boca': '🤭',
    'sorpresa discreta': '🫢',
    'mirando entre manos': '🫣',
    'silencio': '🤫',
    'pensando': '🤔',
    'saludo militar': '🫡',
    'boca cerrada': '🤐',
    'sospechoso': '🤨',
    'neutral': '😐',
    'sin expresion': '😑',
    'sin boca': '😶',
    'cara disuelta': '🫥',
    'hombre': '♂️',
    'mujer': '♀️',
    'genero neutro': '⚧️',
    'agotado': '😓',
    'triste': '😔',
    'dormido': '😴',
    'enfermo': '😷',
    'fiebre': '🤒',
    'dolor de cabeza': '🤕',
    'náuseas': '🤢',
    'vomitando': '🤮',
    'resfriado': '🤧',
    'calor': '🥵',
    'frío': '🥶',
    'mareado': '🥴',
    'aturdido': '😵',
    'explosión mental': '🤯',
    'con gafas': '😎',
    'analítico': '🧐',
    'confundido': '😕',
    'preocupado': '😟',
    'llorando': '😢',
    'llanto fuerte': '😭',
    'enojado': '😡',
    'molesto': '😠',
    'muy enojado': '🤬',
    'diablillo': '👿',
    'calavera': '💀',
    'peligro': '☠️',
    'fantasma': '👻',
    'extraterrestre': '👽',
    'robot': '🤖',
    'bebé': '👶',
    'niño': '🧒',
    'niño varón': '👦',
    'niña': '👧',
    'persona': '🧑',
    'hombre': '👨',
    'mujer': '👩',
    'anciano': '👴',
    'anciana': '👵',
    'levantando la mano': '🙋',
    'celebración': '🙌',
    'aplauso': '👏',
    'apretón de manos': '🤝',
    'oración': '🙏',
    'fuerza': '💪',
    'audífono': '🦻',
    'inteligencia': '🧠',
    'corazón anatómico': '🫀',
    'ojos': '👀',
    'oreja': '👂',
    'nariz': '👃',
    'boca': '👄',
    'lengua': '👅',
    'diente': '🦷',
    'hueso': '🦴',
    'pierna biónica': '🦿',
    'brazo biónico': '🦾',
    'discapacidad': '♿',
    'caminando': '🚶',
    'de pie': '🧍',
    'de rodillas': '🧎',
    'persona en silla de ruedas manual': '🧑‍🦽',
    'persona en silla de ruedas motorizada': '🧑‍🦼',
    'persona con bastón blanco': '🧑‍🦯',
     'profesor': '🧑‍🏫',
    'maestro': '👨‍🏫',
    'maestra': '👩‍🏫',
    'estudiante': '🧑‍🎓',
    'alumno': '👨‍🎓',
    'alumna': '👩‍🎓',
    'libros': '📚',
    'libro abierto': '📖',
    'escribiendo': '📝',
    'cuaderno': '📓',
    'cuaderno de anillas': '📒',
    'calendario': '📅',
    'agenda': '📆',
    'gráfica de barras': '📊',
    'gráfica ascendente': '📈',
    'gráfica descendente': '📉',
    'clip': '📎',
    'chincheta': '📌',
    'ubicación': '📍',
    'carpeta': '📁',
    'carpeta abierta': '📂',
    'archivo': '🗂️',
    'bolígrafo': '🖊️',
    'pluma': '🖋️',
    'lápiz': '✏️',
    'crayón': '🖍️',
    'pincel': '🖌️',
    'tijeras': '✂️',
    'regla': '📏',
    'escuadra': '📐',
    'lupa': '🔍',
    'lupa con texto': '🔎',
    'microscopio': '🔬',
    'telescopio': '🔭',
    'tubo de ensayo': '🧪',
    'placa de petri': '🧫',
    'adn': '🧬',
    'matraz': '⚗️',
    'marcador': '🔖',
    'escuela': '🏫',
    'edificio': '🏢',
    'ordenador': '🖥️',
    'computadora portátil': '💻',
    'impresora': '🖨️',
    'teclado': '⌨️',
    'ratón': '🖱️',
    'programador': '🧑‍💻',
    'hombre programador': '👨‍💻',
    'mujer programadora': '👩‍💻',
    'antena parabólica': '📡',
    'internet': '🌐',
    'enlace': '🔗',
}



def texto_para_accesibilidad(texto, modo="html"):
    doc = nlp(texto)
    texto_adaptado = []
    for token in doc:
        palabra = token.text.lower()
        if palabra in PICTOGRAMAS:
            texto_adaptado.append(f"{token.text} {PICTOGRAMAS[palabra]}")
        else:
            texto_adaptado.append(token.text)
    texto_adaptado = " ".join(texto_adaptado)

    if modo == "html":
        # Generar HTML adaptado para personas con discapacidad
        texto_html = f"""
        <div style="
            font-family: 'OpenDyslexic', Arial, sans-serif;
            font-size: 1.2em;
            line-height: 1.8;
            color: #222;
            background-color: #f9f9f9;
            padding: 20px;
            border-radius: 10px;
            text-align: left;
            word-spacing: 0.2em;
        ">
            {texto_adaptado}
        </div>
        """
        return texto_html
    